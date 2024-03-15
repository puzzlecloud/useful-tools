import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def txt_to_word(input_txt, output_doc):
    with open(input_txt, 'r') as file:
        lines = file.read().split('\n\n')

    # 创建一个新的Word文档
    doc = Document()

    # 设置字体样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # 对输入进行排序
    lines.sort(key=lambda x: (x.split('\n')[1], x.split('\n')[2]))

    for line in lines:
        parts = line.split('\n')
        author, initial, year, title, url = parts[1], parts[2], parts[3], parts[4], parts[5]

        # 检查是否需要斜体
        italic = False
        if title.startswith('#x'):
            italic = True
            title = title[2:].strip()

        # 创建一个新的段落并设置其样式
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 添加作者、年份和标题信息
        para.add_run(f'{author}, {initial} {year}, ')
        para.add_run(f'{title} ').italic = italic
        para.add_run(f'viewed {datetime.datetime.now().strftime("%d %B %Y")}, ')
        para.add_run(f'<{url}>.')

    # 保存Word文档
    doc.save(output_doc)

# 使用函数
txt_to_word('input.txt', 'output.docx')